import streamlit as st
import google.generativeai as genai
import json
import re
import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Page Config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="PrivacyScope AI",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── Custom CSS ───────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

.stApp {
    background: linear-gradient(135deg, #0B1120 0%, #0F172A 50%, #0B1120 100%);
}
[data-testid="stSidebar"] {
    background: rgba(15,23,42,0.97) !important;
    border-right: 1px solid rgba(255,255,255,0.06);
}
[data-testid="stSidebar"] * { color: #CBD5E1 !important; }

.hero-card {
    background: linear-gradient(135deg,rgba(14,165,233,0.09),rgba(139,92,246,0.09));
    border: 1px solid rgba(14,165,233,0.22);
    border-radius: 20px;
    padding: 36px;
    margin-bottom: 28px;
    text-align: center;
}
.info-card {
    background: rgba(255,255,255,0.04);
    border: 1px solid rgba(255,255,255,0.09);
    border-radius: 14px;
    padding: 20px 22px;
    margin-bottom: 14px;
}
.section-bar {
    background: rgba(14,165,233,0.07);
    border-left: 4px solid #0EA5E9;
    border-radius: 0 12px 12px 0;
    padding: 14px 20px;
    margin-bottom: 18px;
}
.free-badge {
    display:inline-block;
    background:rgba(16,185,129,0.15);
    color:#10B981;
    border:1px solid rgba(16,185,129,0.3);
    border-radius:100px;
    padding:3px 12px;
    font-size:12px;
    font-weight:600;
}
.stButton > button {
    background: linear-gradient(135deg,#0EA5E9,#8B5CF6) !important;
    color: white !important;
    border: none !important;
    border-radius: 10px !important;
    font-weight: 600 !important;
    transition: all .2s !important;
}
.stButton > button:hover { opacity:.9 !important; transform:translateY(-1px) !important; }
.stTextInput > div > div > input,
.stTextArea  > div > div > textarea {
    background: rgba(255,255,255,0.06) !important;
    border: 1px solid rgba(255,255,255,0.12) !important;
    border-radius: 10px !important;
    color: #F1F5F9 !important;
}
.stTextInput label,.stSelectbox label,.stMultiSelect label,
.stRadio label,.stCheckbox label { color:#94A3B8 !important; font-size:13px !important; font-weight:500 !important; }
.stTabs [data-baseweb="tab-list"] {
    background:rgba(255,255,255,0.04); border-radius:12px; padding:4px; gap:4px;
}
.stTabs [data-baseweb="tab"]       { border-radius:8px !important; color:#94A3B8 !important; }
.stTabs [aria-selected="true"]     { background:rgba(14,165,233,0.2) !important; color:#0EA5E9 !important; }
[data-testid="metric-container"]   { background:rgba(255,255,255,0.04); border:1px solid rgba(255,255,255,0.08); border-radius:12px; padding:16px; }
.streamlit-expanderHeader          { background:rgba(255,255,255,0.03) !important; border:1px solid rgba(255,255,255,0.07) !important; border-radius:10px !important; color:#CBD5E1 !important; }
hr { border-color:rgba(255,255,255,0.06) !important; }
h1,h2,h3 { color:#F1F5F9 !important; }
p,li      { color:#94A3B8; }
</style>
""", unsafe_allow_html=True)

# ─── Questionnaire Sections ────────────────────────────────────────────────────
SECTIONS = [
    {
        "id": "org_overview", "title": "Organisational Overview", "icon": "🏢", "color": "#0EA5E9",
        "questions": [
            {"id":"q1","text":"Are there any subsidiaries, affiliates, or joint ventures to be included in this engagement?","type":"yesno","ai_key":"has_subsidiaries","elaboration":True},
            {"id":"q2","text":"Is there a centralised Cybersecurity/IT, HR and Legal team responsible for supporting all business functions?","type":"yesno","ai_key":"has_centralized_teams"},
            {"id":"q3","text":"What is the approximate employee strength?","type":"single_choice","ai_key":"employee_count",
             "options":["<500","500 – 1,000","1,000 – 5,000",">5,000"]},
        ],
    },
    {
        "id": "governance", "title": "Governance & Accountability", "icon": "⚖️", "color": "#8B5CF6",
        "questions": [
            {"id":"q4","text":"Has a Privacy Governance Committee or Privacy Office been set up?","type":"single_choice","ai_key":"privacy_governance",
             "options":["Yes, centralised global office","Yes, regional offices","No, decisions taken by IT/Legal/Other","No formal structure"]},
            {"id":"q5","text":"Who takes decisions on personal data usage?","type":"multi_choice","ai_key":"decision_makers",
             "options":["Privacy Office","Legal & Compliance","IT Security","Business Unit Heads","Other"]},
            {"id":"q6","text":"What is the current status of the organisation's privacy policy framework?","type":"single_choice","ai_key":"policy_status",
             "options":["Existing framework in place (requires update)","Drafted but not implemented","Needs to be formulated from scratch","Other"]},
        ],
    },
    {
        "id": "business_lines", "title": "Business Lines & Stakeholders", "icon": "📊", "color": "#10B981",
        "questions": [
            {"id":"q7","text":"Which are the core business lines of the organisation?","type":"dynamic_multi","ai_key":"business_lines"},
            {"id":"q8","text":"Which internal teams or function heads are key stakeholders for processing personal data?","type":"dynamic_multi","ai_key":"stakeholder_teams"},
        ],
    },
    {
        "id": "data_ecosystem", "title": "Data Ecosystem", "icon": "🖥️", "color": "#F59E0B",
        "questions": [
            {"id":"q9","text":"List all customer-facing interfaces used by the organisation.","type":"dynamic_multi","ai_key":"customer_interfaces"},
            {"id":"q10","text":"Which core systems/applications handle personal data?","type":"dynamic_multi","ai_key":"core_systems"},
            {"id":"q11","text":"Do you use any data discovery or mapping tools internally?","type":"yesno","ai_key":"data_discovery_tools","elaboration":True},
            {"id":"q12","text":"Where is personal data stored and hosted?","type":"multi_choice","ai_key":"data_storage",
             "options":["On-premise","Cloud","Hybrid","Third-party hosted"]},
        ],
    },
    {
        "id": "data_subjects", "title": "Data Subjects & Data Types", "icon": "👥", "color": "#EF4444",
        "questions": [
            {"id":"q13","text":"Which categories of individuals (\"Data Subjects\") does the organisation process personal data for?","type":"dynamic_multi","ai_key":"data_subjects"},
            {"id":"q14","text":"What types of personal data does the organisation collect, store, or process?","type":"dynamic_multi","ai_key":"data_types"},
        ],
    },
]

# ─── Gemini Analysis ───────────────────────────────────────────────────────────
PROMPT = """You are a senior privacy consultant doing a pre-scoping privacy assessment.

Analyse the company: "{org_name}"
Website: {website}

Research this company thoroughly. Based on its industry, size and public information, determine:
- Main business lines / products / services
- Industry/sector
- Approximate headcount
- Key internal departments
- Technology systems (ERP, CRM, HRIS, payroll, etc.)
- Types of customers/data subjects
- Types of personal data processed
- Customer-facing digital channels

Return ONLY a JSON object — no markdown fences, no explanation, just raw JSON:
{{
  "company_summary": "2-3 sentence description of what this company does",
  "sector": "e.g. Manufacturing / Banking / Healthcare / IT Services",
  "has_subsidiaries": true,
  "has_centralized_teams": true,
  "employee_count": "1,000 – 5,000",
  "privacy_governance": "No formal structure",
  "decision_makers": ["Legal & Compliance", "IT Security"],
  "policy_status": "Drafted but not implemented",
  "business_lines": [
    {{"name": "Specific business line", "selected": true, "reason": "why relevant to privacy"}},
    {{"name": "Another business line", "selected": true, "reason": "reason"}}
  ],
  "stakeholder_teams": [
    {{"name": "HR & People Operations", "selected": true, "reason": "processes employee data"}},
    {{"name": "IT & Cybersecurity",     "selected": true, "reason": "manages systems"}},
    {{"name": "Legal & Compliance",     "selected": true, "reason": "data governance"}}
  ],
  "customer_interfaces": [
    {{"name": "Web Portal",    "selected": true,  "reason": "customer self-service"}},
    {{"name": "Mobile App",    "selected": false, "reason": "not confirmed"}},
    {{"name": "Call Centre",   "selected": true,  "reason": "customer support"}}
  ],
  "core_systems": [
    {{"name": "SAP ERP",          "selected": true,  "reason": "enterprise resource planning"}},
    {{"name": "Salesforce CRM",   "selected": true,  "reason": "customer management"}},
    {{"name": "Workday HRIS",     "selected": false, "reason": "possible but unconfirmed"}}
  ],
  "data_discovery_tools": false,
  "data_storage": ["On-premise", "Cloud"],
  "data_subjects": [
    {{"name": "Employees",              "selected": true, "reason": "all companies have employees"}},
    {{"name": "Customers / End-users",  "selected": true, "reason": "core business"}},
    {{"name": "Vendors & Contractors",  "selected": true, "reason": "supply chain"}}
  ],
  "data_types": [
    {{"name": "Identity Data (Name, ID proofs)",  "selected": true, "reason": "KYC / onboarding"}},
    {{"name": "Contact & Demographic Data",       "selected": true, "reason": "all stakeholders"}},
    {{"name": "Financial Data",                   "selected": true, "reason": "payroll / transactions"}},
    {{"name": "Employment Data",                  "selected": true, "reason": "HR operations"}}
  ],
  "confidence": "high",
  "notes": "Key privacy observations for the consultant"
}}

Make the business_lines, stakeholder_teams, core_systems, data_subjects and data_types SPECIFIC to this company and its sector — not generic.
Include 4-8 items per dynamic array. For employee_count use exactly one of: "<500", "500 – 1,000", "1,000 – 5,000", ">5,000".
"""


def analyse_company(org_name: str, website: str, api_key: str) -> dict:
    genai.configure(api_key=api_key)

    prompt_text = PROMPT.format(
        org_name=org_name,
        website=website.strip() if website.strip() else "not provided — use your knowledge",
    )

    # Model fallback chain — try each until one works
    # gemini-1.5-flash has the best free-tier quota (15 RPM, 1M TPM)
    models_to_try = [
        "gemini-1.5-flash",
        "gemini-1.5-flash-8b",
        "gemini-1.5-pro",
        "gemini-2.0-flash",
    ]

    last_error = None
    for model_name in models_to_try:
        try:
            model = genai.GenerativeModel(model_name=model_name)
            response = model.generate_content(
                prompt_text,
                generation_config={"temperature": 0.3, "max_output_tokens": 4096},
            )
            text = response.text.strip()

            # Strip markdown fences if present
            text = re.sub(r"^```(?:json)?", "", text).strip()
            text = re.sub(r"```$",           "", text).strip()

            # Extract JSON object
            match = re.search(r"\{[\s\S]*\}", text)
            if match:
                return json.loads(match.group(0))
        except Exception as e:
            last_error = e
            # If quota error, try next model; otherwise re-raise
            err_str = str(e).lower()
            if any(k in err_str for k in ["429", "quota", "resource_exhausted", "rate"]):
                continue
            raise e

    # All models exhausted
    raise ValueError(
        f"All Gemini models hit quota limits. Please wait a minute and try again.\n\n"
        f"Tip: The free tier allows ~15 requests/minute. Last error: {last_error}"
    )


# ─── Word Document Export ─────────────────────────────────────────────────────
def hex_to_rgb(hex_str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def shade_cell(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.lstrip("#"))
    tcPr.append(shd)

def para_bottom_border(para, color_hex):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b    = OxmlElement("w:bottom")
    b.set(qn("w:val"),   "single")
    b.set(qn("w:sz"),    "6")
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), color_hex.lstrip("#"))
    pBdr.append(b)
    pPr.append(pBdr)

def generate_docx(org_name: str, analysis: dict, answers: dict) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── Title block ──
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("PRE-SCOPING PRIVACY QUESTIONNAIRE")
    tr.bold = True; tr.font.size = Pt(22)
    tr.font.color.rgb = hex_to_rgb("#0EA5E9")
    tr.font.name = "Calibri"

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run("Generated by PrivacyScope AI  ·  Protiviti India Member Firm  ·  Data Privacy Team")
    sr.font.size = Pt(9); sr.font.color.rgb = hex_to_rgb("#64748B")

    doc.add_paragraph()

    # ── Meta table ──
    mt = doc.add_table(rows=4, cols=2)
    mt.style = "Table Grid"
    meta = [
        ("Organisation", org_name),
        ("Date",         datetime.now().strftime("%d %B %Y")),
        ("Sector",       analysis.get("sector","—")),
        ("AI Confidence",analysis.get("confidence","medium").capitalize()),
    ]
    mcols = ["0EA5E9","8B5CF6","10B981","F59E0B"]
    for i,(lbl,val) in enumerate(meta):
        row = mt.rows[i]
        shade_cell(row.cells[0], mcols[i])
        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(lbl); lr.bold=True
        lr.font.color.rgb = RGBColor(255,255,255); lr.font.size=Pt(10)
        row.cells[1].paragraphs[0].add_run(val).font.size = Pt(10)

    doc.add_paragraph()

    # ── Company summary ──
    ch = doc.add_paragraph()
    cr = ch.add_run("COMPANY SUMMARY")
    cr.bold=True; cr.font.size=Pt(12); cr.font.color.rgb=hex_to_rgb("#0EA5E9")
    para_bottom_border(ch, "#0EA5E9")

    doc.add_paragraph(analysis.get("company_summary","")).runs[0].font.size = Pt(10)
    if analysis.get("notes"):
        np_ = doc.add_paragraph()
        nr  = np_.add_run("📝 Consultant Note: " + analysis["notes"])
        nr.italic=True; nr.font.size=Pt(9); nr.font.color.rgb=hex_to_rgb("#F59E0B")
    doc.add_paragraph()

    # ── Sections ──
    sec_colors = {
        "org_overview":"0EA5E9","governance":"8B5CF6",
        "business_lines":"10B981","data_ecosystem":"F59E0B","data_subjects":"EF4444",
    }

    for sec in SECTIONS:
        col = sec_colors.get(sec["id"],"0EA5E9")

        sh = doc.add_paragraph()
        sr_ = sh.add_run(f"{sec['icon']}  {sec['title'].upper()}")
        sr_.bold=True; sr_.font.size=Pt(13)
        sr_.font.color.rgb=hex_to_rgb(f"#{col}"); sr_.font.name="Calibri"
        sh.paragraph_format.space_before = Pt(10)
        para_bottom_border(sh, f"#{col}")

        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        tbl.columns[0].width = Cm(0.9)
        tbl.columns[1].width = Cm(8.5)
        tbl.columns[2].width = Cm(7.3)

        hdr = tbl.rows[0]
        for ci,ht in enumerate(["S.N","Attributes","Response"]):
            shade_cell(hdr.cells[ci], col)
            hp = hdr.cells[ci].paragraphs[0]
            hr_ = hp.add_run(ht)
            hr_.bold=True; hr_.font.color.rgb=RGBColor(255,255,255); hr_.font.size=Pt(9)

        for qi,q in enumerate(sec["questions"]):
            row = tbl.add_row()
            row.cells[0].paragraphs[0].add_run(str(qi+1)).font.size=Pt(9)
            row.cells[1].paragraphs[0].add_run(q["text"]).font.size=Pt(9)

            ans  = answers.get(q["id"])
            ap   = row.cells[2].paragraphs[0]

            def tick(t):
                r=ap.add_run("☑  "+t+"\n"); r.font.size=Pt(9)
                r.font.color.rgb=RGBColor(5,150,105)
            def cross(t):
                r=ap.add_run("☐  "+t+"\n"); r.font.size=Pt(9)
                r.font.color.rgb=RGBColor(100,116,139)

            if ans is None:
                ap.add_run("—").font.size=Pt(9)
            elif q["type"]=="yesno":
                tick("Yes") if ans else cross("Yes")
                cross("No") if ans else tick("No")
            elif q["type"]=="single_choice":
                for opt in q["options"]:
                    tick(opt) if opt==ans else cross(opt)
            elif q["type"]=="multi_choice":
                cur = ans if isinstance(ans,list) else []
                for opt in q["options"]:
                    tick(opt) if opt in cur else cross(opt)
            elif q["type"]=="dynamic_multi":
                items = ans if isinstance(ans,list) else []
                if items:
                    for item in items:
                        nm = item if isinstance(item,str) else item.get("name","")
                        tick(nm)
                else:
                    ap.add_run("—").font.size=Pt(9)

        doc.add_paragraph()

    # ── Footer ──
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        f"Auto-generated by PrivacyScope AI on {datetime.now().strftime('%d %B %Y')}. "
        "Please review all responses before use."
    )
    fr.italic=True; fr.font.size=Pt(8); fr.font.color.rgb=hex_to_rgb("#64748B")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─── Session State ────────────────────────────────────────────────────────────
for k,v in {"phase":"landing","analysis":None,"answers":{},"org_name":"","website":"","gemini_key":""}.items():
    if k not in st.session_state:
        st.session_state[k] = v

# Try key from Streamlit secrets (for hosted version)
if not st.session_state.gemini_key:
    st.session_state.gemini_key = os.environ.get("GEMINI_API_KEY","") or st.secrets.get("GEMINI_API_KEY","") if hasattr(st,"secrets") else os.environ.get("GEMINI_API_KEY","")

# ─── Sidebar ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 🛡️ PrivacyScope AI")
    st.markdown("*Pre-Scoping Privacy Questionnaire*")
    st.markdown("<span class='free-badge'>✓ Powered by Google Gemini — Free</span>", unsafe_allow_html=True)
    st.divider()

    key_input = st.text_input(
        "🔑 Google Gemini API Key",
        type="password",
        value=st.session_state.gemini_key,
        placeholder="AIza...",
        help="Free key from aistudio.google.com — no credit card needed",
    )
    if key_input:
        st.session_state.gemini_key = key_input

    st.markdown("""
    <div style='background:rgba(16,185,129,0.08);border:1px solid rgba(16,185,129,0.2);
         border-radius:10px;padding:12px;margin-top:8px;font-size:12px;color:#94A3B8;line-height:1.7'>
    <b style='color:#10B981'>Get your FREE key:</b><br>
    1. Go to <b>aistudio.google.com</b><br>
    2. Sign in with Google<br>
    3. Click <b>"Get API Key"</b><br>
    4. Paste it above ⬆️
    </div>
    """, unsafe_allow_html=True)

    st.divider()

    if st.session_state.phase == "results" and st.session_state.analysis:
        st.markdown("### 📋 Sections")
        for i,sec in enumerate(SECTIONS):
            total = len(sec["questions"])
            done  = sum(1 for q in sec["questions"] if st.session_state.answers.get(q["id"]) is not None)
            st.markdown(f"**{sec['icon']} {sec['title']}**")
            st.progress(done/total if total else 0, text=f"{done}/{total} answered")

        st.divider()
        a = st.session_state.analysis
        st.markdown(f"**Sector:** {a.get('sector','')}")
        st.markdown(f"**Confidence:** {a.get('confidence','').upper()}")
        st.markdown(f"<small style='color:#64748B'>{a.get('company_summary','')}</small>", unsafe_allow_html=True)

        st.divider()
        if st.button("🔄 New Analysis", use_container_width=True):
            st.session_state.update({"phase":"landing","analysis":None,"answers":{}})
            st.rerun()


# ─── LANDING ──────────────────────────────────────────────────────────────────
if st.session_state.phase == "landing":

    st.markdown("""
    <div class='hero-card'>
        <div style='font-size:52px;margin-bottom:10px'>🛡️</div>
        <h1 style='color:#F1F5F9;font-size:34px;font-weight:800;margin:0 0 8px'>
            PrivacyScope AI
        </h1>
        <p style='color:#94A3B8;font-size:16px;margin:0 0 14px'>
            Intelligent Pre-Scoping Privacy Questionnaire Generator
        </p>
        <span class='free-badge'>✓ 100% Free · Powered by Google Gemini</span>
    </div>
    """, unsafe_allow_html=True)

    c1,c2,c3,c4 = st.columns(4)
    for col,icon,lbl in zip([c1,c2,c3,c4],
        ["🌐","🎯","🖥️","📄"],
        ["Reads company\nwebsite","Identifies business\nlines & sector","Maps IT systems\n& interfaces","Exports Word\ndocument (.docx)"]):
        with col:
            st.markdown(f"""
            <div class='info-card' style='text-align:center'>
                <div style='font-size:30px'>{icon}</div>
                <div style='font-size:12px;color:#94A3B8;margin-top:6px'>{lbl}</div>
            </div>""", unsafe_allow_html=True)

    st.markdown("---")

    left, right = st.columns([1.3, 1])

    with left:
        st.markdown("### 🏢 Enter Organisation Details")
        org  = st.text_input("Organisation Name *", placeholder="e.g. Infosys Limited", value=st.session_state.org_name)
        site = st.text_input("Company Website (recommended)", placeholder="e.g. https://www.infosys.com", value=st.session_state.website,
                             help="Providing the website makes AI results significantly more accurate.")

        if not st.session_state.gemini_key:
            st.warning("⚠️ Enter your free Gemini API key in the sidebar to continue.")

        clicked = st.button("⚡  Analyse & Generate Questionnaire", use_container_width=True, type="primary")

        if clicked:
            if not org.strip():
                st.error("Please enter the organisation name.")
            elif not st.session_state.gemini_key:
                st.error("Please enter your Gemini API Key in the sidebar.")
            else:
                st.session_state.org_name = org
                st.session_state.website  = site

                steps = [
                    ("🌐","Searching company website & public sources…"),
                    ("🏢","Identifying business lines & sector…"),
                    ("🖥️","Mapping technology systems…"),
                    ("👥","Profiling data subjects & data types…"),
                    ("⚖️","Assessing governance posture…"),
                    ("✅","Building tailored questionnaire…"),
                ]

                box = st.empty()
                import time

                def render_steps(active):
                    rows = ""
                    for i,(ic,lb) in enumerate(steps):
                        if i < active:
                            clr,sym = "#10B981","✓"
                        elif i == active:
                            clr,sym = "#0EA5E9","⟳"
                        else:
                            clr,sym = "#475569","○"
                        rows += f"<div style='color:{clr};font-size:14px;margin:8px 0;font-weight:600'>{ic} {sym} {lb}</div>"
                    pct = int((active+1)/len(steps)*100)
                    box.markdown(f"""
                    <div class='info-card'>
                      <div style='color:#F1F5F9;font-weight:700;font-size:16px;margin-bottom:14px'>
                        🤖 Analysing: {org}
                      </div>
                      {rows}
                      <div style='margin-top:16px;background:rgba(255,255,255,0.08);border-radius:100px;height:6px'>
                        <div style='width:{pct}%;height:100%;background:linear-gradient(90deg,#0EA5E9,#8B5CF6);border-radius:100px;transition:width .4s'></div>
                      </div>
                    </div>""", unsafe_allow_html=True)

                for i in range(len(steps)-1):
                    render_steps(i); time.sleep(0.7)
                render_steps(len(steps)-1)

                try:
                    result = analyse_company(org, site, st.session_state.gemini_key)
                    init_ans = {}
                    for sec in SECTIONS:
                        for q in sec["questions"]:
                            val = result.get(q["ai_key"])
                            if val is not None:
                                init_ans[q["id"]] = val
                    st.session_state.analysis = result
                    st.session_state.answers  = init_ans
                    st.session_state.phase    = "results"
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Analysis failed: {e}")

    with right:
        st.markdown("### 💡 What gets auto-identified")
        st.markdown("""
        <div class='info-card'>
        <div style='font-size:13px;color:#94A3B8;line-height:2.1'>
        ✅ <b style='color:#F1F5F9'>Business lines</b> specific to your sector<br>
        ✅ <b style='color:#F1F5F9'>IT systems</b> — ERP, CRM, HRIS, etc.<br>
        ✅ <b style='color:#F1F5F9'>Data subjects</b> — employees, customers…<br>
        ✅ <b style='color:#F1F5F9'>Data types</b> — identity, financial, health…<br>
        ✅ <b style='color:#F1F5F9'>Stakeholder teams</b> mapped to your org<br>
        ✅ <b style='color:#F1F5F9'>Customer interfaces</b> — app, portal, etc.<br>
        ✅ <b style='color:#F1F5F9'>Governance posture</b> assessment<br>
        ✅ <b style='color:#F1F5F9'>Word document</b> export (.docx)<br>
        </div>
        </div>""", unsafe_allow_html=True)

        st.markdown("### 🎯 Example")
        st.markdown("""
        <div class='info-card'>
        <div style='font-size:12px;color:#94A3B8;line-height:1.9'>
        <b style='color:#0EA5E9'>Input:</b> "HDFC Bank" + website<br><br>
        <b style='color:#10B981'>Auto-identified:</b><br>
        • Business lines: Retail, Corporate, Insurance…<br>
        • Systems: Finacle, Salesforce, Workday…<br>
        • Data subjects: Account holders, KYC…<br>
        • Interfaces: Mobile app, NetBanking…
        </div>
        </div>""", unsafe_allow_html=True)


# ─── RESULTS ──────────────────────────────────────────────────────────────────
elif st.session_state.phase == "results":
    analysis = st.session_state.analysis
    answers  = st.session_state.answers
    org_name = st.session_state.org_name

    # ── Top bar ──
    h1,h2,h3 = st.columns([3,1,1])
    with h1:
        st.markdown(f"""
        <div style='display:flex;align-items:center;gap:14px;margin-bottom:4px'>
          <span style='font-size:28px'>🛡️</span>
          <div>
            <div style='font-size:22px;font-weight:800;color:#F1F5F9'>{org_name}</div>
            <div style='font-size:12px;color:#64748B'>{analysis.get('sector','')} · Pre-Scoping Privacy Questionnaire</div>
          </div>
        </div>""", unsafe_allow_html=True)
    with h2:
        total_q = sum(len(s["questions"]) for s in SECTIONS)
        done_q  = sum(1 for s in SECTIONS for q in s["questions"] if answers.get(q["id"]) is not None)
        st.metric("Progress", f"{done_q}/{total_q}", "questions answered")
    with h3:
        docx_bytes = generate_docx(org_name, analysis, answers)
        st.download_button(
            "⬇️  Export .docx", data=docx_bytes,
            file_name=f"Privacy_Questionnaire_{org_name.replace(' ','_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    st.divider()

    # ── Tabs ──
    tabs = st.tabs([f"{s['icon']} {s['title']}" for s in SECTIONS])

    for tab, sec in zip(tabs, SECTIONS):
        with tab:
            st.markdown(f"""
            <div class='section-bar' style='border-left-color:{sec["color"]}'>
              <span style='font-size:22px'>{sec["icon"]}</span>
              <span style='font-size:17px;font-weight:700;color:{sec["color"]};margin-left:10px'>{sec["title"]}</span>
              <span style='color:#64748B;font-size:13px;margin-left:10px'>— Review AI-suggested answers</span>
            </div>""", unsafe_allow_html=True)

            for qi,q in enumerate(sec["questions"]):
                with st.expander(f"**Q{qi+1}.** {q['text']}", expanded=True):
                    current = answers.get(q["id"])

                    if q["type"] == "yesno":
                        default_i = 0 if current is True else 1
                        val = st.radio("Select:", ["Yes","No"], index=default_i, horizontal=True, key=f"yn_{q['id']}")
                        answers[q["id"]] = (val == "Yes")
                        if q.get("elaboration") and val == "Yes":
                            st.text_area("Please elaborate:", key=f"el_{q['id']}", height=70)

                    elif q["type"] == "single_choice":
                        opts = q["options"]
                        idx  = opts.index(current) if current in opts else 0
                        answers[q["id"]] = st.radio("Select one:", opts, index=idx, key=f"sc_{q['id']}")

                    elif q["type"] == "multi_choice":
                        cur = current if isinstance(current,list) else []
                        sel = []
                        cols = st.columns(min(len(q["options"]),3))
                        for oi,opt in enumerate(q["options"]):
                            with cols[oi % len(cols)]:
                                if st.checkbox(opt, value=opt in cur, key=f"mc_{q['id']}_{oi}"):
                                    sel.append(opt)
                        answers[q["id"]] = sel

                    elif q["type"] == "dynamic_multi":
                        ai_items = analysis.get(q["ai_key"], [])
                        cur_names = {(i if isinstance(i,str) else i.get("name",""))
                                     for i in (current or [])}

                        st.markdown(
                            "<div style='font-size:11px;color:#0EA5E9;font-weight:600;margin-bottom:8px'>"
                            "🤖 AI-identified options — tick to include in questionnaire</div>",
                            unsafe_allow_html=True)

                        selected = []
                        cols2 = st.columns(2)
                        for oi,item in enumerate(ai_items):
                            nm  = item.get("name","") if isinstance(item,dict) else item
                            rsn = item.get("reason","") if isinstance(item,dict) else ""
                            with cols2[oi % 2]:
                                if st.checkbox(nm, value=nm in cur_names, key=f"dm_{q['id']}_{oi}",
                                               help=rsn or None):
                                    selected.append({"name":nm,"reason":rsn})

                        custom = st.text_input("➕ Add custom:", key=f"cust_{q['id']}", placeholder="Type and press Enter…")
                        if custom:
                            selected.append({"name":custom,"reason":"Manually added"})
                        answers[q["id"]] = selected

            st.session_state.answers = answers

    # ── Bottom export ──
    st.divider()
    _,_,btn_col = st.columns([2,1,1])
    with btn_col:
        st.download_button(
            "⬇️  Export Final Questionnaire",
            data=generate_docx(org_name, analysis, answers),
            file_name=f"Privacy_Questionnaire_{org_name.replace(' ','_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
